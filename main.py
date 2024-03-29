import copy
import json
import os
import random
from math import ceil

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import Pt

import pandas as pd

import subprocess

import argparse


def configure_document():
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    return document


def generate_master_table(master_answers, filename):
    df = pd.DataFrame.from_dict(master_answers).T
    df.to_csv(filename)


def generate_document(data: dict, model_number: int, filename: str, input_folder) -> dict:
    correct_options = {}

    document = configure_document()  # type: Document

    header = data['header'].format(model_number)
    pre, post = header.split('<b>')
    p = document.add_paragraph(pre)
    pre, post = post.split('</b>')
    p.add_run(pre).bold = True
    p.add_run(post)

    random.seed(None)
    questions = copy.deepcopy(data['questions'])
    random.shuffle(questions)

    n_cols = min(6, len(questions))
    n_rows = 2 * ceil(len(questions) / n_cols)

    document.add_paragraph(
        'Ao fim da prova, preencha a tabela abaixo, assinalando '
        'qual alternativa foi escolhida para cada uma das questões.'
    )

    table = document.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # single lines in all cells

    counter = 1
    for i in range(n_rows):
        if i % 2 == 0:
            for j in range(n_cols):
                if counter > len(questions):
                    break
                table.cell(i, j).text = 'Questão {0}'.format(counter)
                table.cell(i, j).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(i, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                counter += 1
        else:
            for j in range(n_cols):
                table.cell(i, j).text = ' '

    document.add_paragraph('')

    for i, question in enumerate(questions):
        correct_value = question['options'][question['correct_index']]

        p = document.add_paragraph('')

        p.add_run(question['name'].format(i + 1)).bold = True
        p.add_run(question['description'])

        if question['image'] is not None:
            width = eval(question['image']['width'])  # type: Cm

            document.add_picture(os.path.join(input_folder, question['image']['path']), width=width)
            p = document.paragraphs[-1]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        options = question['options']
        random.shuffle(options)

        index = options.index(correct_value)
        correct_options['Questão {0}'.format(i + 1)] = chr(ord('a') + index)

        p = document.add_paragraph('{0}) {1}\n'.format('a', options[0]))
        op = ord('b')
        for option in options[1:]:
            p.add_run('{0}) {1}\n'.format(chr(op), option))
            op += 1

    # document.add_page_break()

    document.save(filename)

    return correct_options


def main(input_file, output_folder, n_assignments=1):
    filename = os.path.basename(input_file).split('.')[0]
    input_folder = os.sep.join(input_file.split(os.sep)[:-1])

    with open(input_file, 'r', encoding='utf-8') as datafile:
        data = json.load(datafile)

    master_answers = {}
    for i in range(1, n_assignments + 1):
        output_docx = '{0}_model_{1:02d}.{2}'.format(filename, i, 'docx')

        correct_answers = generate_document(
            data, i, os.path.join(output_folder, output_docx), input_folder=input_folder
        )

        script_path = os.path.dirname(os.path.abspath(__file__))

        subprocess.call([
            'soffice', '--headless',
            '--convert-to', 'pdf', os.path.join(script_path, output_folder, output_docx),
            '--outdir', os.path.join(script_path, output_folder)
        ])

        master_answers['Modelo {0:02d}'.format(i)] = correct_answers

    generate_master_table(master_answers, os.path.join(output_folder, '{0}_master_table.csv'.format(filename)))


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Automatically generates exames based on a JSON file. Randomizes question and options orders.'
    )

    parser.add_argument(
        '--input-file', action='store', required=True,
        help='Path to input JSON file.'
    )
    parser.add_argument(
        '--output-folder', action='store', required=True,
        help='Path to output folder where exams will be stored, both in docx, pdf and csv format.'
    )
    parser.add_argument(
        '--n-assignments', action='store', required=False, type=int, default=1,
        help='Number of different exams to generate. Defaults to 1'
    )

    args = parser.parse_args()

    main(input_file=args.input_file, output_folder=args.output_folder, n_assignments=args.n_assignments)
